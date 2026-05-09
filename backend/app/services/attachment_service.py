import base64
import io
import re
import uuid
from csv import reader as csv_reader
from pathlib import Path

from fastapi import HTTPException, UploadFile
from openai import OpenAI
from docx import Document
from pypdf import PdfReader
from sqlalchemy.orm import Session

from app.ai.llm import DEFAULT_LITELLM_PROXY_URL, resolve_model_name
from app.core.config import settings
from app.repositories import attachment_repo


PROJECT_ROOT = Path(__file__).resolve().parents[3]
BACKEND_ROOT = Path(__file__).resolve().parents[2]


# Allowed file types and their MIME types
ALLOWED_FILE_TYPES = {
    "image": {
        "extensions": {".jpg", ".jpeg", ".png", ".gif", ".webp", ".svg"},
        "mimes": {"image/jpeg", "image/png", "image/gif", "image/webp", "image/svg+xml"},
    },
    "video": {
        "extensions": {".mp4", ".webm", ".mov", ".avi"},
        "mimes": {"video/mp4", "video/webm", "video/quicktime", "video/x-msvideo"},
    },
    "code": {
        "extensions": {".py", ".js", ".ts", ".tsx", ".jsx", ".java", ".cpp", ".c", ".go", ".rs", ".rb", ".php", ".sql", ".html", ".css", ".json", ".yaml", ".xml", ".txt"},
        "mimes": {"text/plain", "text/x-python", "text/javascript", "text/typescript", "application/json", "text/yaml", "text/xml", "text/html", "text/css"},
    },
    "document": {
        "extensions": {".pdf", ".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx", ".txt", ".csv"},
        "mimes": {"application/pdf", "application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "application/vnd.ms-powerpoint", "application/vnd.openxmlformats-officedocument.presentationml.presentation", "text/plain", "text/csv"},
    },
    "formula": {
        "extensions": {".tex", ".txt", ".md"},
        "mimes": {"text/plain", "text/markdown"},
    },
}


def get_upload_directory() -> Path:
    """Get or create upload directory"""
    upload_dir = settings.upload_dir or "./uploads"
    path = Path(upload_dir)
    if not path.is_absolute():
        path = (PROJECT_ROOT / path).resolve()
    path.mkdir(parents=True, exist_ok=True)
    return path


def validate_file_type(file_name: str, mime_type: str, file_type: str) -> bool:
    """Validate if file type and extension are allowed"""
    if file_type not in ALLOWED_FILE_TYPES:
        return False

    allowed = ALLOWED_FILE_TYPES[file_type]
    file_ext = Path(file_name).suffix.lower()

    # Check extension and MIME type
    return file_ext in allowed["extensions"] and mime_type in allowed["mimes"]


def validate_file_size(file_size: int) -> bool:
    """Validate file size against configured limit"""
    max_size_bytes = settings.max_upload_mb * 1024 * 1024
    return file_size <= max_size_bytes


def generate_storage_path(message_id: uuid.UUID, file_name: str) -> str:
    """Generate unique storage path for uploaded file"""
    # Create subdirectory based on message ID for organization
    file_ext = Path(file_name).suffix
    unique_filename = f"{uuid.uuid4()}{file_ext}"
    storage_subdir = str(message_id)[:8]  # Use first 8 chars of UUID for grouping
    storage_path = f"{storage_subdir}/{unique_filename}"
    return storage_path


async def save_uploaded_file(file: UploadFile, storage_path: str) -> None:
    """Save uploaded file to storage"""
    upload_dir = get_upload_directory()
    file_full_path = upload_dir / storage_path

    # Ensure directory exists
    file_full_path.parent.mkdir(parents=True, exist_ok=True)

    try:
        # Save file
        with open(file_full_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to save file: {str(e)}")


def get_file_path(storage_path: str) -> Path:
    """Get full file path for a stored attachment"""
    upload_dir = get_upload_directory()
    primary = upload_dir / storage_path
    if primary.exists():
        return primary

    # Backward-compatible fallback for files saved when backend was started from backend/ cwd.
    legacy_dir = BACKEND_ROOT / (settings.upload_dir or "./uploads")
    legacy_path = legacy_dir / storage_path
    if legacy_path.exists():
        return legacy_path

    return primary


def file_exists(storage_path: str) -> bool:
    """Check if file exists"""
    return get_file_path(storage_path).exists()


def delete_uploaded_file(storage_path: str) -> bool:
    """Delete uploaded file from storage"""
    try:
        file_path = get_file_path(storage_path)
        if file_path.exists():
            file_path.unlink()
            # Try to clean up empty directories
            try:
                file_path.parent.rmdir()
            except OSError:
                pass
            return True
        return False
    except Exception:
        return False


async def process_and_store_attachment(
    file: UploadFile,
    message_id: uuid.UUID,
    file_type: str,
    db: Session,
) -> dict:
    """Process, validate, and store attachment"""

    # Get file content to check size
    content = await file.read()
    await file.seek(0)  # Reset for saving

    file_size = len(content)
    mime_type = file.content_type or "application/octet-stream"

    # Validation checks
    if not validate_file_size(file_size):
        raise HTTPException(
            status_code=413,
            detail=f"File size exceeds limit of {settings.max_upload_mb}MB",
        )

    if not validate_file_type(file.filename, mime_type, file_type):
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed for {file_type} attachment. Check file extension and MIME type.",
        )

    # Generate storage path and save file
    storage_path = generate_storage_path(message_id, file.filename)
    await save_uploaded_file(file, storage_path)

    # Create attachment record in database
    attachment = attachment_repo.create_attachment(
        db=db,
        message_id=message_id,
        file_name=file.filename,
        file_type=file_type,
        mime_type=mime_type,
        file_size=file_size,
        storage_path=storage_path,
    )

    return {
        "id": attachment.id,
        "file_name": attachment.file_name,
        "file_type": attachment.file_type,
        "mime_type": attachment.mime_type,
        "file_size": attachment.file_size,
        "storage_path": attachment.storage_path,
        "upload_timestamp": attachment.upload_timestamp,
    }


async def delete_attachment_with_file(db: Session, attachment_id: uuid.UUID) -> bool:
    """Delete attachment record and associated file"""
    attachment = attachment_repo.get_attachment(db, attachment_id)
    if not attachment:
        return False

    # Delete file from storage
    delete_uploaded_file(attachment.storage_path)

    # Delete record from database
    return attachment_repo.delete_attachment(db, attachment_id)


def _read_text_file(file_path: Path) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return file_path.read_text(errors="ignore")


def _extract_from_pdf(file_path: Path) -> str:
    pdf = PdfReader(str(file_path))
    blocks: list[str] = []
    for page in pdf.pages:
        text = page.extract_text() or ""
        if text.strip():
            blocks.append(text.strip())
    return "\n\n".join(blocks)


def _extract_from_docx(file_path: Path) -> str:
    doc = Document(str(file_path))
    lines: list[str] = []

    for p in doc.paragraphs:
        text = (p.text or "").strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = "\n".join(
                    (para.text or "").strip() for para in cell.paragraphs if (para.text or "").strip()
                ).strip()
                cells.append(cell_text)
            row_text = " | ".join(c for c in cells if c)
            if row_text:
                lines.append(row_text)

    # Keep ordering while removing duplicates from merged table cells.
    unique_lines = list(dict.fromkeys(lines))
    return "\n".join(unique_lines)


def _extract_from_xlsx(file_path: Path) -> str:
    try:
        import pandas as pd
    except Exception as exc:
        raise HTTPException(status_code=500, detail="pandas is required for spreadsheet extraction") from exc

    xls = pd.ExcelFile(file_path)
    parts: list[str] = []
    for sheet in xls.sheet_names[:5]:
        frame = xls.parse(sheet)
        if frame.empty:
            continue
        parts.append(f"[Sheet: {sheet}]")
        parts.append(frame.fillna("").to_csv(index=False).strip())
    return "\n\n".join(parts)


def _extract_from_csv(file_path: Path) -> str:
    text = _read_text_file(file_path)
    sio = io.StringIO(text)
    rows = list(csv_reader(sio))[:200]
    return "\n".join(",".join(cell.strip() for cell in row) for row in rows)


def _extract_from_non_image_file(attachment, file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if attachment.file_type == "video":
        return _extract_from_video_with_llm(file_path, attachment.mime_type or "video/mp4")

    if attachment.file_type in {"code", "formula"}:
        return _read_text_file(file_path)

    if attachment.file_type != "document":
        raise HTTPException(status_code=400, detail=f"Unsupported attachment type for extraction: {attachment.file_type}")

    if suffix in {".txt", ".md", ".json", ".xml", ".yaml", ".yml"}:
        return _read_text_file(file_path)
    if suffix == ".csv":
        return _extract_from_csv(file_path)
    if suffix == ".pdf":
        return _extract_from_pdf(file_path)
    if suffix == ".docx":
        return _extract_from_docx(file_path)
    if suffix == ".xlsx":
        return _extract_from_xlsx(file_path)

    raise HTTPException(
        status_code=400,
        detail=f"Extraction is not supported for this document format: {suffix or attachment.file_name}",
    )


def _get_llm_client_and_model() -> tuple[OpenAI, str]:
    api_key = settings.litellm_api_key
    if not api_key:
        raise HTTPException(status_code=500, detail="Missing LITELLM_API_KEY in backend configuration")

    proxy_url = (settings.litellm_proxy_url or DEFAULT_LITELLM_PROXY_URL).rstrip("/")
    base_url = proxy_url if proxy_url.endswith("/v1") else f"{proxy_url}/v1"
    client = OpenAI(api_key=api_key, base_url=base_url)
    model_name = resolve_model_name()
    return client, model_name


def _extract_from_video_with_llm(file_path: Path, mime_type: str) -> str:
    try:
        import cv2
    except Exception as exc:
        raise HTTPException(status_code=500, detail="opencv-python-headless is required for video extraction") from exc

    cap = cv2.VideoCapture(str(file_path))
    if not cap.isOpened():
        raise HTTPException(status_code=400, detail="Unable to read uploaded video")

    frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT) or 0)
    max_frames = 6

    if frame_count > 1:
        indices = sorted({int(i * (frame_count - 1) / (max_frames - 1)) for i in range(max_frames)})
    else:
        indices = [0]

    encoded_frames: list[str] = []
    for idx in indices:
        cap.set(cv2.CAP_PROP_POS_FRAMES, idx)
        ok, frame = cap.read()
        if not ok or frame is None:
            continue
        ok_jpg, buffer = cv2.imencode(".jpg", frame)
        if not ok_jpg:
            continue
        encoded_frames.append(base64.b64encode(buffer.tobytes()).decode("utf-8"))

    cap.release()

    if not encoded_frames:
        raise HTTPException(status_code=400, detail="Could not sample frames from uploaded video")

    client, model_name = _get_llm_client_and_model()
    content_parts: list[dict] = [
        {
            "type": "text",
            "text": (
                "Analyze these sampled video frames. "
                "1) Extract any visible text exactly as written. "
                "2) Briefly summarize what the video appears to show. "
                "Return plain text only with headings: Visible Text and Summary."
            ),
        }
    ]
    for frame_b64 in encoded_frames:
        content_parts.append(
            {
                "type": "image_url",
                "image_url": {"url": f"data:image/jpeg;base64,{frame_b64}"},
            }
        )

    response = client.chat.completions.create(
        model=model_name,
        temperature=0,
        messages=[{"role": "user", "content": content_parts}],
    )
    content = response.choices[0].message.content if response.choices else ""
    return (content or "").strip()


def extract_text_from_image_attachment(db: Session, attachment_id: uuid.UUID) -> str:
    """Extract visible text from an uploaded image attachment using the configured multimodal LLM."""
    attachment = attachment_repo.get_attachment(db, attachment_id)
    if not attachment:
        raise HTTPException(status_code=404, detail="Attachment not found")

    file_path = get_file_path(attachment.storage_path)
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Attachment file not found")

    if attachment.file_type != "image":
        return _extract_from_non_image_file(attachment, file_path)

    mime_type = attachment.mime_type or "image/jpeg"
    image_bytes = file_path.read_bytes()
    b64 = base64.b64encode(image_bytes).decode("utf-8")
    client, model_name = _get_llm_client_and_model()

    response = client.chat.completions.create(
        model=model_name,
        temperature=0,
        messages=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "text",
                        "text": (
                            "Extract all visible text from this image exactly as written. "
                            "Return plain text only. Keep line breaks. "
                            "If no readable text exists, return an empty string."
                        ),
                    },
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{mime_type};base64,{b64}",
                        },
                    },
                ],
            }
        ],
    )

    content = response.choices[0].message.content if response.choices else ""
    return (content or "").strip()


def _fallback_summarize_text(extracted_text: str, attachment_type: str, file_name: str) -> str:
    text = (extracted_text or "").strip()
    if not text:
        return "No readable content was extracted from this file."

    compact = re.sub(r"\s+", " ", text).strip()
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    sentence_candidates = [seg.strip() for seg in re.split(r"(?<=[.!?])\s+", compact) if seg.strip()]
    key_points = sentence_candidates[:5] if sentence_candidates else lines[:5]

    if not key_points and compact:
        key_points = [compact[:220]]

    important_data: list[str] = []
    for line in lines:
        if any(ch.isdigit() for ch in line) or "@" in line or "$" in line or "%" in line:
            important_data.append(line)
        if len(important_data) >= 5:
            break

    short_overview = compact[:260].strip()
    if len(compact) > 260:
        short_overview += "..."

    summary_lines = [
        "Overview",
        f"This {attachment_type} file ({file_name}) contains information about: {short_overview}",
        "",
        "Key Points",
    ]
    summary_lines.extend([f"- {point[:220]}" for point in key_points])

    summary_lines.extend(["", "Important Data"])
    if important_data:
        summary_lines.extend([f"- {item[:220]}" for item in important_data])
    else:
        summary_lines.append("- No specific numeric or structured values were identified.")

    summary_lines.extend(
        [
            "",
            "Suggested Next Steps",
            "- Ask for a deeper section-wise summary if you need more detail.",
            "- Ask for action items or decisions based on this file.",
        ]
    )

    return "\n".join(summary_lines).strip()


def summarize_attachment_content(extracted_text: str, attachment_type: str, file_name: str) -> str:
    """Create a readable summary for extracted attachment content."""
    text = (extracted_text or "").strip()
    if not text:
        return "No readable content was extracted from this file."

    # Keep prompt payload bounded for large files.
    bounded_text = text[:12000]

    if settings.litellm_api_key:
        try:
            client, model_name = _get_llm_client_and_model()
            response = client.chat.completions.create(
                model=model_name,
                temperature=0.2,
                messages=[
                    {
                        "role": "system",
                        "content": (
                            "You are a precise analyst. Summarize extracted file content clearly and neatly. "
                            "Return plain text only with these headings exactly: "
                            "Overview, Key Points, Important Data, Suggested Next Steps."
                        ),
                    },
                    {
                        "role": "user",
                        "content": (
                            f"File name: {file_name}\n"
                            f"Attachment type: {attachment_type}\n\n"
                            "Extracted content:\n"
                            f"{bounded_text}"
                        ),
                    },
                ],
            )
            content = response.choices[0].message.content if response.choices else ""
            summarized = (content or "").strip()
            if summarized:
                # Avoid returning near-verbatim extraction output.
                norm_source = re.sub(r"\s+", " ", bounded_text).strip().lower()
                norm_summary = re.sub(r"\s+", " ", summarized).strip().lower()
                if len(norm_source) > 80:
                    if norm_summary in norm_source:
                        return _fallback_summarize_text(text, attachment_type, file_name)
                    overlap_prefix = norm_source[: min(400, len(norm_source))]
                    if overlap_prefix and overlap_prefix in norm_summary and len(norm_summary) > 0.8 * len(norm_source):
                        return _fallback_summarize_text(text, attachment_type, file_name)
                return summarized
        except Exception:
            # Fall back to deterministic local summarization if LLM summary fails.
            pass

    return _fallback_summarize_text(text, attachment_type, file_name)
